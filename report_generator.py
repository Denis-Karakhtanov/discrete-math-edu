from docx import Document
import openpyxl
from datetime import datetime
import os

class ReportGenerator:
    def __init__(self):
        """Инициализация генератора отчетов"""
        self.reports_dir = "reports"
        if not os.path.exists(self.reports_dir):
            os.makedirs(self.reports_dir)
    
    def export_to_docx(self, user_login, score, weak_topics):
        """Экспорт отчета в .docx"""
        try:
            doc = Document()
            doc.add_heading(f"Отчет по результатам теста для {user_login}", 0)
            doc.add_paragraph(f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Результат: {score:.2f}%")
            
            if weak_topics:
                doc.add_heading("Слабые темы:", level=1)
                for topic in weak_topics:
                    doc.add_paragraph(f"- {topic}")
            
            file_path = os.path.join(self.reports_dir, f"report_{user_login}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            doc.save(file_path)
            return file_path
        except Exception as e:
            raise Exception(f"Ошибка при создании .docx отчета: {str(e)}")
    
    def export_to_xlsx(self, user_login, score, weak_topics):
        """Экспорт отчета в .xlsx"""
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Отчет"
            
            ws["A1"] = "Пользователь"
            ws["B1"] = user_login
            ws["A2"] = "Дата"
            ws["B2"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            ws["A3"] = "Результат"
            ws["B3"] = f"{score:.2f}%"
            
            if weak_topics:
                ws["A5"] = "Слабые темы"
                for i, topic in enumerate(weak_topics, start=6):
                    ws[f"A{i}"] = topic
            
            file_path = os.path.join(self.reports_dir, f"report_{user_login}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
            wb.save(file_path)
            return file_path
        except Exception as e:
            raise Exception(f"Ошибка при создании .xlsx отчета: {str(e)}")